import logging
import json
import os # For joining paths and checking file existence
from typing import List, Dict, Optional, Any
from rank_bm25 import BM25Okapi # For BM25 indexing

from config import settings
from core.llm_service import LLMService, LLMServiceError
from core.embedding_service import EmbeddingService, EmbeddingServiceError
from core.reranker_service import RerankerService, RerankerServiceError
from core.document_processor import DocumentProcessor, DocumentProcessorError
from core.vector_store import VectorStore, VectorStoreError

from agents.topic_analyzer_agent import TopicAnalyzerAgent, TopicAnalyzerAgentError
from agents.outline_generator_agent import OutlineGeneratorAgent, OutlineGeneratorAgentError
from agents.content_retriever_agent import ContentRetrieverAgent, ContentRetrieverAgentError
from agents.chapter_writer_agent import ChapterWriterAgent, ChapterWriterAgentError
from agents.evaluator_agent import EvaluatorAgent, EvaluatorAgentError
from agents.refiner_agent import RefinerAgent, RefinerAgentError
from agents.report_compiler_agent import ReportCompilerAgent, ReportCompilerAgentError

logger = logging.getLogger(__name__)

class ReportGenerationPipelineError(Exception):
    """Custom exception for pipeline errors."""
    pass

class ReportGenerationPipeline:
    """
    Orchestrates the entire report generation process, from document processing
    to final report compilation, using a sequence of specialized agents.
    Now supports parent-child chunking, multiple document types, and hybrid retrieval.
    """

    def __init__(self,
                 llm_service: LLMService,
                 embedding_service: EmbeddingService,
                 reranker_service: Optional[RerankerService] = None,
                 # Document processing params (can be overridden by CLI in main.py)
                 parent_chunk_size: int = settings.DEFAULT_PARENT_CHUNK_SIZE,
                 parent_chunk_overlap: int = settings.DEFAULT_PARENT_CHUNK_OVERLAP,
                 child_chunk_size: int = settings.DEFAULT_CHILD_CHUNK_SIZE,
                 child_chunk_overlap: int = settings.DEFAULT_CHILD_CHUNK_OVERLAP,
                 # Retrieval params (can be overridden by CLI in main.py)
                 vector_top_k: int = settings.DEFAULT_VECTOR_STORE_TOP_K,
                 keyword_top_k: int = settings.DEFAULT_KEYWORD_SEARCH_TOP_K,
                 hybrid_alpha: float = settings.DEFAULT_HYBRID_SEARCH_ALPHA,
                 final_top_n_retrieval: Optional[int] = None, # top_n for ContentRetriever output
                 # Pipeline execution params
                 max_refinement_iterations: int = settings.DEFAULT_MAX_REFINEMENT_ITERATIONS
                ):
        """
        Initializes the ReportGenerationPipeline.
        Params related to chunking and retrieval can be passed here, often from main.py CLI args.
        """
        self.llm_service = llm_service
        self.embedding_service = embedding_service
        self.reranker_service = reranker_service
        self.max_refinement_iterations = max_refinement_iterations

        # Initialize services and agents, passing down relevant parameters
        self.document_processor = DocumentProcessor(
            parent_chunk_size=parent_chunk_size,
            parent_chunk_overlap=parent_chunk_overlap,
            child_chunk_size=child_chunk_size,
            child_chunk_overlap=child_chunk_overlap,
            supported_extensions=settings.SUPPORTED_DOC_EXTENSIONS
        )
        self.vector_store = VectorStore(embedding_service=self.embedding_service)

        # BM25 index will be built after documents are processed
        self.bm25_index: Optional[BM25Okapi] = None
        self.all_child_chunks_for_bm25_corpus: List[Dict[str, Any]] = []


        self.topic_analyzer = TopicAnalyzerAgent(llm_service=self.llm_service)
        self.outline_generator = OutlineGeneratorAgent(llm_service=self.llm_service)

        # ContentRetrieverAgent will be initialized later, after bm25_index is ready
        self.content_retriever: Optional[ContentRetrieverAgent] = None
        self.retrieval_params = {
            "default_vector_top_k": vector_top_k,
            "default_keyword_top_k": keyword_top_k,
            "default_hybrid_alpha": hybrid_alpha,
            "default_final_top_n": final_top_n_retrieval or vector_top_k # Default final_top_n to vector_top_k if not specified
        }

        self.chapter_writer = ChapterWriterAgent(llm_service=self.llm_service)
        self.evaluator = EvaluatorAgent(llm_service=self.llm_service)
        self.refiner = RefinerAgent(llm_service=self.llm_service)
        self.report_compiler = ReportCompilerAgent(add_table_of_contents=True)

        logger.info("ReportGenerationPipeline initialized.")

    def _initialize_content_retriever(self):
        """Initializes the ContentRetrieverAgent once BM25 index is available."""
        if not self.content_retriever:
            self.content_retriever = ContentRetrieverAgent(
                vector_store=self.vector_store,
                bm25_index=self.bm25_index,
                all_child_chunks_for_bm25=self.all_child_chunks_for_bm25_corpus,
                reranker_service=self.reranker_service,
                **self.retrieval_params
            )
            logger.info("ContentRetrieverAgent initialized with BM25 index.")


    def _process_and_load_data(self, data_path: str):
        """
        Scans a directory for supported files, extracts text, performs parent-child
        chunking, loads child chunks into VectorStore, and builds BM25 index.
        """
        logger.info(f"Starting document processing from data_path: {data_path}")

        if not os.path.isdir(data_path):
            logger.error(f"Data path '{data_path}' is not a valid directory.")
            raise ReportGenerationPipelineError(f"Invalid data_path: {data_path} is not a directory.")

        all_parent_child_data: List[Dict[str, Any]] = []
        processed_file_count = 0

        for filename in os.listdir(data_path):
            file_path = os.path.join(data_path, filename)
            if not os.path.isfile(file_path):
                logger.debug(f"Skipping non-file item: {filename}")
                continue

            _, extension = os.path.splitext(filename.lower())
            if extension not in settings.SUPPORTED_DOC_EXTENSIONS:
                logger.debug(f"Skipping unsupported file type: {filename} (extension: {extension})")
                continue

            try:
                logger.info(f"Processing file: {file_path}")
                raw_text = self.document_processor.extract_text_from_file(file_path)
                if not raw_text.strip():
                    logger.warning(f"No text extracted from file: {file_path}. Skipping.")
                    continue

                # Use filename (without extension) or a UUID as document ID
                doc_id_base = os.path.splitext(filename)[0]
                parent_child_chunks = self.document_processor.split_text_into_parent_child_chunks(
                    raw_text, source_document_id=doc_id_base
                )
                all_parent_child_data.extend(parent_child_chunks)
                logger.info(f"Processed {file_path} into {len(parent_child_chunks)} parent chunks.")
                processed_file_count +=1

            except (DocumentProcessorError, FileNotFoundError) as e:
                logger.error(f"Failed to process file {file_path}: {e}. It will be skipped.")
            except Exception as e: # Catch any other unexpected error during individual file processing
                logger.error(f"Unexpected error processing file {file_path}: {e}. Skipping.", exc_info=True)


        if not all_parent_child_data:
            logger.error("No parent-child chunks were generated from any file in the data_path. Cannot proceed.")
            raise ReportGenerationPipelineError("No usable content extracted or chunked from provided data_path.")

        logger.info(f"Total parent chunks from {processed_file_count} files: {len(all_parent_child_data)}")

        # Load into VectorStore (embeds child chunks and stores parent/child metadata)
        try:
            logger.info(f"Adding data to VectorStore...")
            self.vector_store.add_documents(all_parent_child_data)
            logger.info("Successfully added document data to VectorStore.")
        except VectorStoreError as e:
            logger.error(f"Failed to add documents to VectorStore: {e}")
            raise ReportGenerationPipelineError(f"VectorStore population failed: {e}")

        # Build BM25 Index using child chunks from the vector_store's document_store
        # self.vector_store.document_store should contain {'child_id', 'child_text', 'parent_id', 'parent_text', ...}
        self.all_child_chunks_for_bm25_corpus = [
            {"child_id": item['child_id'], "child_text": item['child_text']}
            for item in self.vector_store.document_store
        ]

        if self.all_child_chunks_for_bm25_corpus:
            tokenized_corpus_for_bm25 = [
                self.content_retriever._tokenize_query(item['child_text']) if self.content_retriever else item['child_text'].lower().split() # Use agent's tokenizer if available
                for item in self.all_child_chunks_for_bm25_corpus
            ]
            self.bm25_index = BM25Okapi(tokenized_corpus_for_bm25)
            logger.info(f"BM25 index built successfully with {len(tokenized_corpus_for_bm25)} child chunks.")
        else:
            logger.warning("No child chunks available to build BM25 index.")
            self.bm25_index = None # Ensure it's None if no corpus

        # Now that BM25 index is ready (or None), initialize ContentRetrieverAgent
        self._initialize_content_retriever()


    def _parse_outline_to_chapter_titles(self, markdown_outline: str) -> List[str]:
        # (This method remains largely the same as before, but ensure it's robust)
        titles = []
        lines = markdown_outline.strip().split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue

            # Heuristic for Markdown list items or headers
            if line.startswith(("- ", "* ", "+ ")):
                title = line[2:].strip()
                if title: titles.append(title)
            elif line.startswith("#"):
                title = line.lstrip("# ").strip()
                if title: titles.append(title)

        if not titles:
            logger.warning("Could not parse any chapter titles from the generated outline.")
        logger.debug(f"Parsed outline titles for content generation: {titles}")
        return titles

    def run(self, user_topic: str, data_path: str, report_title: Optional[str] = None) -> str:
        """
        Executes the full report generation pipeline.

        Args:
            user_topic (str): The main topic for the report.
            data_path (str): Path to the directory containing source documents.
            report_title (Optional[str]): Desired title for the final report.

        Returns:
            str: The final compiled report in Markdown format.
        """
        logger.info(f"Starting report generation pipeline for topic: '{user_topic}' using data from '{data_path}'")
        final_report_title = report_title or f"关于“{user_topic}”的分析报告"

        # 1. Process and Load Documents (and build BM25 index)
        try:
            self._process_and_load_data(data_path)
        except ReportGenerationPipelineError as e:
             raise

        if not self.content_retriever: # Should have been initialized in _process_and_load_data
            logger.error("ContentRetrieverAgent was not initialized. This indicates an issue in the setup phase.")
            raise ReportGenerationPipelineError("ContentRetrieverAgent failed to initialize.")


        # 2. Analyze Topic
        logger.info("Step 2: Analyzing topic...")
        analyzed_topic_details = self.topic_analyzer.run(user_topic)
        logger.info(f"Topic analysis complete. Generalized CN topic: {analyzed_topic_details.get('generalized_topic_cn')}")

        # 3. Generate Outline
        logger.info("Step 3: Generating outline...")
        markdown_outline = self.outline_generator.run(analyzed_topic_details)
        if not markdown_outline.strip():
            raise ReportGenerationPipelineError("Outline generation failed: Empty outline.")
        logger.info("Outline generation complete.")

        # 4. Process each chapter
        logger.info("Step 4: Processing chapters...")
        chapter_titles_for_writing = self._parse_outline_to_chapter_titles(markdown_outline)
        if not chapter_titles_for_writing:
            raise ReportGenerationPipelineError("Failed to parse chapter titles from outline.")

        compiled_chapter_contents: Dict[str, str] = {}
        for i, chapter_title in enumerate(chapter_titles_for_writing):
            logger.info(f"Processing chapter {i+1}/{len(chapter_titles_for_writing)}: '{chapter_title}'")

            query_keywords_cn = analyzed_topic_details.get('keywords_cn', [])
            # query_keywords_en = analyzed_topic_details.get('keywords_en', []) # Not used in simple concat below
            retrieval_query = f"{chapter_title} {' '.join(query_keywords_cn)}".strip() # Focus on Chinese keywords for query

            retrieved_parent_chunks_data = self.content_retriever.run(query=retrieval_query) # Uses agent's defaults

            # retrieved_parent_chunks_data is now List[Dict{'document': parent_text, 'score':..., other_meta...}]
            logger.info(f"Retrieved {len(retrieved_parent_chunks_data)} parent chunks for chapter '{chapter_title}'.")
            if not retrieved_parent_chunks_data:
                logger.warning(f"No relevant parent chunks found for chapter '{chapter_title}'. Content might be sparse.")

            current_chapter_content = self.chapter_writer.run(
                chapter_title=chapter_title,
                retrieved_content=retrieved_parent_chunks_data # Pass the list of dicts
            )
            logger.info(f"Initial draft for '{chapter_title}'. Length: {len(current_chapter_content)}")

            if current_chapter_content.strip() and not current_chapter_content.startswith("Error"):
                for ref_iter in range(self.max_refinement_iterations):
                    logger.info(f"Refinement iteration {ref_iter + 1}/{self.max_refinement_iterations} for '{chapter_title}'")
                    evaluation = self.evaluator.run(content_to_evaluate=current_chapter_content)
                    logger.info(f"Evaluation for '{chapter_title}': Score {evaluation.get('score')}")

                    refined_content = self.refiner.run(
                        original_content=current_chapter_content,
                        evaluation_feedback=evaluation
                    )
                    logger.info(f"Refinement for '{chapter_title}'. Orig len: {len(current_chapter_content)}, New len: {len(refined_content)}")
                    current_chapter_content = refined_content
            else:
                logger.warning(f"Skipping refinement for '{chapter_title}' due to empty or error content.")
            compiled_chapter_contents[chapter_title] = current_chapter_content

        # 5. Compile Report
        logger.info("Step 5: Compiling final report...")
        final_report_md = self.report_compiler.run(
            report_title=final_report_title,
            markdown_outline=markdown_outline,
            chapter_contents=compiled_chapter_contents,
            report_topic_details=analyzed_topic_details
        )
        logger.info("Report compilation complete.")

        logger.info("Report generation pipeline finished successfully.")
        return final_report_md


if __name__ == '__main__':
    logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    logger.info("ReportGenerationPipeline (Hybrid & Parent-Child) Example Start")

    # --- Mock Services (similar to previous main.py example, adapted for new flow) ---
    class MockLLMServiceForPipeline(LLMService): # Copied from previous main.py, should be fine
        def __init__(self): super().__init__(api_url="mock://llm", model_name="mock-llm")
        def chat(self, query, system_prompt, **kwargs):
            logger.debug(f"MockLLMService.chat called. Query starts with: {query[:60]}")
            if "你是一个主题分析专家" in query:
                return json.dumps({"generalized_topic_cn": "模拟混合检索主题", "generalized_topic_en": "Mock Hybrid Topic", "keywords_cn": ["父子分块", "关键词"], "keywords_en": ["parent-child", "keyword"]})
            if "你是一个报告大纲撰写助手" in query:
                return "- 章节一：父子分块策略\n  - 1.1 定义\n- 章节二：混合检索实现"
            if "你是一位专业的报告撰写员" in query:
                title_search = "章节标题：\n"; ts_idx = query.find(title_search); te_idx = query.find("\n", ts_idx + len(title_search) if ts_idx != -1 else -1)
                title = query[ts_idx + len(title_search):te_idx].strip() if ts_idx != -1 and te_idx != -1 else "未知章节"
                return f"这是“{title}”的模拟章节内容，基于父块上下文撰写。"
            if "你是一位资深的报告评审员" in query: return json.dumps({"score": 82, "feedback_cn": "模拟反馈：良好。", "evaluation_criteria_met": {"relevance": "高", "fluency": "良好", "completeness": "良好", "accuracy": "良好"}})
            if "你是一位报告修改专家" in query: return query.split("原始内容：\n---\n")[1].split("\n---")[0] + "\n（已根据模拟反馈进行修订 - 父子分块版）"
            return "Unknown LLM query for mock."
        def get_model(self, model_name): return self

    class MockEmbeddingServiceForPipeline(EmbeddingService): # Copied
        def __init__(self): super().__init__(api_url="mock://emb", model_name="mock-emb")
        def create_embeddings(self, texts): return [[0.1] * 10 for _ in texts]
        def get_model(self, model_name): return self

    class MockRerankerServiceForPipeline(RerankerService): # Copied
        def __init__(self): super().__init__(api_url="mock://rerank", model_name="mock-rerank")
        def rerank(self, query, documents, top_n=None):
            res = [{"document": d, "relevance_score": 0.9-(i*0.1), "original_index": i} for i,d in enumerate(documents)]
            return res[:top_n] if top_n else res
        def get_model(self, model_name): return self

    # --- Dummy Data Path Setup ---
    dummy_data_dir = "temp_pipeline_data"
    if not os.path.exists(dummy_data_dir): os.makedirs(dummy_data_dir)

    # Create dummy files (TXT, DOCX - PDF is harder to make with text simply)
    with open(os.path.join(dummy_data_dir, "doc1.txt"), "w", encoding="utf-8") as f:
        f.write("这是第一个文本文档。它讨论了苹果和香蕉。\n\n第二段关于橙子。")
    try:
        import docx as python_docx_lib # Alias to avoid conflict with module name
        doc = python_docx_lib.Document()
        doc.add_paragraph("这是一个Word文档，关于红色汽车。")
        doc.add_paragraph("还提到了蓝色自行车和快速交通工具。")
        doc.save(os.path.join(dummy_data_dir, "doc2.docx"))
    except ImportError: logger.warning("python-docx not installed, cannot create dummy .docx for test.")


    try:
        mock_llm = MockLLMServiceForPipeline()
        mock_embed = MockEmbeddingServiceForPipeline()
        mock_rerank = MockRerankerServiceForPipeline()

        pipeline = ReportGenerationPipeline(
            llm_service=mock_llm,
            embedding_service=mock_embed,
            reranker_service=mock_rerank,
            max_refinement_iterations=1
        )

        user_topic_ex = "文档处理与检索新方法"

        logger.info(f"Running pipeline example with topic: '{user_topic_ex}' and data_path: '{dummy_data_dir}'")

        final_report = pipeline.run(
            user_topic=user_topic_ex,
            data_path=dummy_data_dir # Pass directory path
        )

        print("\n" + "="*30 + " FINAL REPORT (Mocked - Hybrid & Parent-Child) " + "="*30)
        print(final_report)
        print("="*80)

    except ReportGenerationPipelineError as e:
        logger.error(f"Pipeline execution failed: {e}", exc_info=True)
    except Exception as e:
        logger.error(f"An unexpected error occurred during pipeline example: {e}", exc_info=True)
    finally:
        import shutil
        if os.path.exists(dummy_data_dir): shutil.rmtree(dummy_data_dir)
        logger.info(f"Cleaned up dummy data directory: {dummy_data_dir}")

    logger.info("ReportGenerationPipeline (Hybrid & Parent-Child) Example End")
